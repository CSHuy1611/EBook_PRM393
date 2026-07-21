from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "KE_HOACH_TEST_VA_UNIT_TEST_MATH_IBOOK.docx"
SKILL_ROOT = Path(
    r"C:\Users\nguye\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F4F6F9"
WHITE = "FFFFFF"
MUTED = "5B6573"
GREEN = "E8F3EC"
AMBER = "FFF4CE"
RED = "FDEBEC"


def set_run_font(run, name="Calibri", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:cantSplit")
    tr_pr.append(elem)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:tblHeader")
    elem.set(qn("w:val"), "true")
    tr_pr.append(elem)


def add_table(doc, headers, rows, weights, *, font_size=8.7, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].height = None
    repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        cant_split(table.rows[-1])
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
            if row_idx % 2 == 1:
                shade_cell(cells[idx], "FAFBFC")
    widths = column_widths_from_weights(weights, 9360)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_field(paragraph):
    paragraph.add_run("Trang ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_numbering_definition(doc, num_id, abstract_id, fmt, text, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def add_list_item(doc, text, *, numbered=False, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42" if numbered else "41")
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_body(doc, text, *, bold_prefix=None, italic=False, color=INK):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color, italic=italic)
    return p


def add_callout(doc, label, text, fill=PALE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "D6DEE8")
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run(label + " ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_test_case(doc, case_id, title, priority, goal, preconditions, steps, expected):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{case_id} - {title} [{priority}]")
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    add_body(doc, f"Mục tiêu: {goal}", bold_prefix="Mục tiêu:")
    add_body(doc, f"Điều kiện: {preconditions}", bold_prefix="Điều kiện:")
    for step in steps:
        add_list_item(doc, step, numbered=True)
    add_body(doc, f"Kết quả mong đợi: {expected}", bold_prefix="Kết quả mong đợi:")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    add_numbering_definition(doc, 41, 41, "bullet", "•")
    add_numbering_definition(doc, 42, 42, "decimal", "%1.")

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("MATH IBOOK  |  TEST PLAN & UNIT TEST PLAN")
    set_run_font(r, size=8, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_field(p)
    for r in p.runs:
        set_run_font(r, size=8, color=MUTED)


def build_document():
    doc = Document()
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("KẾ HOẠCH TEST VÀ UNIT TEST")
    set_run_font(r, size=25, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Luồng chính và các luồng phụ quan trọng - Math IBook")
    set_run_font(r, size=16, bold=False, color=BLUE)

    add_callout(
        doc,
        "Mục tiêu tài liệu:",
        "Biến các rủi ro nghiệp vụ của hệ thống học Toán thành danh mục test có thể triển khai, ưu tiên các invariant về điểm, tiến độ, xu, huy hiệu, phân quyền và đồng bộ offline.",
        LIGHT_BLUE,
    )

    add_table(
        doc,
        ["Thuộc tính", "Giá trị"],
        [
            ("Dự án", "Math IBook - ASP.NET Core 8 + Flutter + PostgreSQL"),
            ("Phiên bản tài liệu", "1.0"),
            ("Ngày lập", "20/07/2026"),
            ("Phạm vi", "Backend service/API, Flutter unit/widget/integration, luồng Student và Admin"),
            ("Căn cứ", "Source hiện tại, README.md, AGENTS.md, TASKS.md và bộ test trong repository"),
            ("Trạng thái", "Kế hoạch đề xuất để triển khai và theo dõi coverage"),
        ],
        [1.55, 4.95],
        font_size=9.2,
        header_fill=LIGHT_GRAY,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("Đối tượng sử dụng: Developer, QA, Reviewer, Product Owner")
    set_run_font(r, size=10, italic=True, color=MUTED)

    doc.add_page_break()

    add_heading(doc, "1. Tóm tắt điều hành", 1)
    add_body(
        doc,
        "Kế hoạch lấy luồng học và làm quiz của Student làm trục chính. Mỗi thay đổi ở luồng này có thể tác động đồng thời tới QuizAttempt, tiến độ bài/chương, số xu, huy hiệu, bảng xếp hạng, dashboard và dữ liệu offline. Vì vậy test không chỉ kiểm tra response đúng mà còn phải chứng minh dữ liệu không bị cộng lặp, không thoái lui trạng thái Passed và luôn giữ BestScore cao nhất.",
    )
    add_body(
        doc,
        "Luồng chính thứ hai là vòng đời nội dung của Admin: tạo chương - bài học - câu hỏi - quiz, validate, publish và hiển thị cho Student. Luồng này bảo vệ chất lượng học liệu và tránh công bố dữ liệu thiếu hoặc sai target.",
    )
    add_callout(
        doc,
        "Ưu tiên cao nhất:",
        "P0 cho xác thực/phân quyền, submit quiz, idempotency ClientAttemptId, reward server-side, giữ pass/best score và offline retry. Các lỗi ở đây có thể làm sai dữ liệu tài khoản hoặc thưởng.",
        AMBER,
    )

    add_heading(doc, "1.1 Kết quả baseline tại thời điểm lập tài liệu", 2)
    add_table(
        doc,
        ["Hạng mục", "Quan sát", "Kết luận sử dụng trong kế hoạch"],
        [
            ("Backend unit test", "27/27 test đạt khi chạy dotnet test ngày 20/07/2026; có 1 warning biến innerEx không dùng.", "Dùng làm regression baseline; chưa thay thế test API/controller."),
            ("Flutter test", "Lần chạy toàn bộ suite dừng do timeout 180 giây, không trả kết quả cuối.", "Cần tách suite/tag và thiết lập timeout CI rõ ràng."),
            ("Ghi nhận trong TASKS.md", "47 test đạt và 10 test legacy lỗi tại MathText/Login ở lần kiểm tra trước.", "Phải triage, cố định test không ổn định trước khi đặt quality gate toàn suite."),
            ("Khoảng trống nổi bật", "Thiếu test Coins, Badges, Profile, OfflineSyncService; thiếu API integration cho JWT, pagination và offline reconcile.", "Đưa vào Phase 1-3 của roadmap."),
        ],
        [1.35, 2.9, 2.25],
    )

    add_heading(doc, "2. Mục tiêu, phạm vi và nguyên tắc", 1)
    add_heading(doc, "2.1 Mục tiêu", 2)
    for item in [
        "Phát hiện regression sớm tại lớp Domain/Application trước khi lỗi lan tới API và UI.",
        "Chứng minh các invariant quan trọng: một attempt chỉ thưởng một lần; Passed không bị hạ; BestScore không giảm; xu/huy hiệu do server quyết định.",
        "Bảo vệ phân quyền Student/Admin, vòng đời refresh token và tài khoản bị khóa.",
        "Bảo vệ luồng offline qua restart, retry, lỗi từng phần và reconcile sau khi có mạng.",
        "Tạo quality gate có thể chạy lặp lại trên máy phát triển và CI.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "2.2 Phạm vi kiểm thử", 2)
    add_table(
        doc,
        ["Trong phạm vi", "Ngoài phạm vi của tài liệu này"],
        [
            ("Unit test Domain/Application C#; controller/API integration; Flutter unit/widget/integration; database constraint và luồng E2E trọng yếu.", "Penetration test chuyên sâu, kiểm thử tải quy mô production, kiểm định sư phạm nội dung toán bởi chuyên gia."),
            ("Student: auth, học bài, quiz, progress, reward, badge, coins, leaderboard, profile, notification, offline sync.", "Kiểm thử thiết bị vật lý đầy đủ theo ma trận OEM; vẫn cần checklist release riêng."),
            ("Admin: nội dung, validation/publish, quiz/question, user status, reward policy, notification, report.", "Migration dữ liệu production hoặc reset database; không được thực hiện trong automated test dùng database thật."),
        ],
        [3.25, 3.25],
    )

    add_heading(doc, "2.3 Nguyên tắc thiết kế test", 2)
    principles = [
        ("Một hành vi - một lý do lỗi", "Test name mô tả Given_When_Then; mỗi test chỉ có một trọng tâm nghiệp vụ."),
        ("Test invariant trước UI", "Logic điểm, thưởng, pass và idempotency phải được khóa ở backend unit/integration."),
        ("Không dùng thời gian ngẫu nhiên", "Inject clock hoặc dùng timestamp cố định; tránh DateTime.Now trong expected."),
        ("Không dùng database chia sẻ", "Integration test tạo schema/container riêng, seed tối thiểu và dọn theo transaction."),
        ("Deterministic offline", "Giữ nguyên ClientAttemptId, mô phỏng timeout/500/400 và kiểm tra queue sau mỗi lần thử."),
        ("Không log dữ liệu nhạy cảm", "Không snapshot password, token, đáp án hoặc body quiz trong output test."),
    ]
    add_table(doc, ["Nguyên tắc", "Áp dụng"], principles, [1.9, 4.6])

    add_heading(doc, "3. Kiến trúc kiểm thử và phân tầng", 1)
    add_body(doc, "Kiến trúc source được ánh xạ thành bốn tầng test. Mỗi tầng có trách nhiệm khác nhau và không nên dùng E2E để thay thế unit test.")
    add_table(
        doc,
        ["Tầng", "Mục đích", "Công cụ/đích chạy", "Tỷ trọng đề xuất"],
        [
            ("Unit - backend", "Domain invariant, service branch, validation, reward/progress/badge.", "xUnit + Moq + FluentAssertions", "50-60% số ca"),
            ("Unit/widget - Flutter", "Parser/model/router, local scoring, DB queue, state loading/error/empty/success.", "flutter_test + fake API/storage", "25-30% số ca"),
            ("API/integration", "Routing, JWT/role, ProblemDetails, EF constraints, transaction/idempotency.", "WebApplicationFactory + PostgreSQL test container", "10-15% số ca"),
            ("E2E", "Xác nhận hành trình thật trên Android/Windows/Web được hỗ trợ.", "integration_test + backend test env", "5-10% số ca"),
        ],
        [1.2, 2.25, 2.05, 1.0],
    )

    add_heading(doc, "3.1 Ma trận ưu tiên", 2)
    add_table(
        doc,
        ["Mức", "Khi nào dùng", "Ví dụ", "Quality gate"],
        [
            ("P0", "Mất dữ liệu, sai tiền/xu, vượt quyền, không đăng nhập/làm bài được.", "JWT role, submit quiz, duplicate reward, offline retry.", "100% đạt trước merge/release"),
            ("P1", "Sai chức năng trọng yếu nhưng có đường tránh hoặc phạm vi hẹp.", "Badge reconcile, profile password, publish validation.", "100% đạt trước release"),
            ("P2", "Sai trạng thái phụ/UX hoặc dữ liệu báo cáo không tức thời.", "Empty/error state, pagination, refresh dashboard.", "Đạt >= 95%; không còn lỗi blocker"),
            ("P3", "Cosmetic/hiếm gặp.", "Spacing, text phụ, animation.", "Theo regression chọn lọc"),
        ],
        [0.6, 2.2, 2.35, 1.35],
    )

    add_heading(doc, "4. Invariant nghiệp vụ bắt buộc khóa bằng test", 1)
    invariants = [
        ("INV-01", "Idempotent attempt", "Cùng UserId + ClientAttemptId trả lại attempt cũ; không tạo QuizAttempt/CoinTransaction/UserBadge trùng.", "P0"),
        ("INV-02", "Best score", "Progress.BestScore chỉ tăng hoặc giữ nguyên, kể cả sync timestamp cũ/mới.", "P0"),
        ("INV-03", "Pass không thoái lui", "Lesson/Chapter đã Passed không trở lại InProgress khi điểm lần sau thấp hơn.", "P0"),
        ("INV-04", "Reward server-side", "Client không thể gửi số xu để cộng; policy, first-pass, retry và daily limit do server tính.", "P0"),
        ("INV-05", "Badge một lần", "Badge đã nhận không được đánh giá/cộng xu/thông báo lại.", "P0"),
        ("INV-06", "Published content only", "Student chỉ thấy/chấm quiz hợp lệ, published, không deleted và đúng target.", "P0"),
        ("INV-07", "Role isolation", "Student không gọi được /api/admin; Admin không đi vào StudentShell ngoài rule định nghĩa.", "P0"),
        ("INV-08", "Offline queue durability", "Attempt còn trong queue khi sync lỗi; retry_count tăng; chỉ mark synced sau response thành công.", "P0"),
        ("INV-09", "User isolation", "Cache/queue phân tách theo user_id; không gửi attempt của tài khoản trước.", "P0"),
        ("INV-10", "Token lifecycle", "Refresh token hết hạn/revoked/thuộc user inactive không phát access token mới.", "P0"),
    ]
    add_table(doc, ["ID", "Invariant", "Điều phải chứng minh", "Ưu tiên"], invariants, [0.75, 1.35, 3.75, 0.65])

    add_heading(doc, "5. Kế hoạch Unit Test backend", 1)
    add_body(doc, "Quy ước triển khai: mock IUnitOfWork/IRepository cho unit test; chỉ dùng database thật ở integration test. Với mỗi service, kiểm tra success path, validation path, not-found/inactive path và branch idempotency.")
    backend_plan = [
        ("AuthService", "Login đúng/sai; inactive; email trùng; refresh hợp lệ/hết hạn/revoked; rotation; logout idempotent; password reset OTP.", "P0", "Mở rộng AuthServiceTests"),
        ("QuizScoringService", "Quiz/target hợp lệ; câu trả lời thiếu/dư/sai ID; score 0/biên/pass/perfect; duplicate ClientAttemptId; lesson/chapter progress; trả correct answer.", "P0", "Mở rộng 2 test hiện có lên ma trận branch"),
        ("QuizRewardService", "Legacy/policy; perfect bonus; retry percent; first pass; chapter completion; daily cap; zero/negative; duplicate transaction.", "P0", "Mở rộng 2 test hiện có"),
        ("ProgressSyncService", "Lesson invalid/unpublished; client timestamp; không tin BestScore client; lấy best từ server attempt; pass không regress; nhiều item.", "P0", "Mở rộng 2 test hiện có"),
        ("BadgeCheckService", "Rule lesson/chapter/coins/passed quiz/streak; legacy fallback; threshold invalid; badge inactive/deleted; award once; reward/notification.", "P0/P1", "Mở rộng 2 test hiện có"),
        ("ContentValidationService", "LaTeX cân bằng; ví dụ lớp 8; exactly one target; 4 options/1 correct; question ngoài chapter; empty content.", "P1", "Giữ 3 test, bổ sung boundary"),
        ("CoinCalculationService", "Biên score, perfect, retry, policy null, clamp non-negative.", "P1", "Mở rộng test hiện có"),
        ("ProfileService", "Get/update/avatar URL/password; trùng email; mật khẩu hiện tại sai; hash mới; user inactive/not found.", "P1", "Mở rộng test hiện có"),
        ("DashboardService", "Empty data; progress aggregation; coins/badges; recent activity; chỉ dữ liệu user.", "P1/P2", "Mở rộng test hiện có"),
        ("QuestionGeneratorService", "lesson/chapter/title/count/level validation; provider error/invalid JSON; đúng số câu và option.", "P1", "Tạo test mới"),
        ("Domain entities", "LessonProgress/ChapterProgress/Quiz target; score ngoài 0-10; unlock/pass idempotent.", "P0", "Duy trì và bổ sung boundary"),
    ]
    add_table(doc, ["Đối tượng", "Nhóm ca kiểm thử", "Ưu tiên", "Hành động"], backend_plan, [1.35, 3.5, 0.75, 1.15], font_size=8.4)

    add_heading(doc, "5.1 Mẫu Arrange - Act - Assert", 2)
    add_test_case(
        doc,
        "UT-BE-QUIZ-01",
        "Retry cùng ClientAttemptId không cộng xu lần hai",
        "P0",
        "Khóa INV-01 và INV-04.",
        "Đã có Student, published quiz, attempt và CoinTransaction có cùng ClientAttemptId.",
        [
            "Arrange repository trả về attempt cũ và reward transaction tương ứng.",
            "Act gọi ScoreQuizAsync với cùng UserId và ClientAttemptId.",
            "Assert kết quả trùng attempt cũ; không Add QuizAttempt/CoinTransaction; Coins không đổi; SaveChanges không tạo bản ghi mới.",
        ],
        "API/service trả kết quả ổn định, không phát sinh thưởng hoặc huy hiệu lặp.",
    )
    add_test_case(
        doc,
        "UT-BE-PROG-02",
        "Điểm thấp sau khi Passed không làm thoái lui tiến độ",
        "P0",
        "Khóa INV-02 và INV-03.",
        "Progress hiện tại Passed, BestScore = 9; attempt mới đạt 4.",
        [
            "Arrange progress và quiz threshold.",
            "Act cập nhật tiến độ từ lần làm mới.",
            "Assert Status vẫn Passed, BestScore vẫn 9, thời gian cập nhật hợp lệ.",
        ],
        "Không mất pass state và best score.",
    )

    add_heading(doc, "6. Kế hoạch Unit/Widget Test Flutter", 1)
    flutter_plan = [
        ("AuthProvider + router", "Restore session, login/logout, role redirect, token refresh failure, inactive/401 chuyển login.", "Unit", "P0"),
        ("ApiClient", "Attach Bearer; refresh một lần; concurrent 401; không loop; timeout/offline mapping; không log body nhạy cảm.", "Unit", "P0"),
        ("OfflineSyncService", "Queue trống; payload; giữ ClientAttemptId; success mark; 400/500/timeout giữ queue; retry summary; per-user isolation.", "Unit", "P0"),
        ("LocalDbService", "Migration; unique client_attempt_id; upsert progress; pending/failed query; retry_count; user_id filter.", "DB integration", "P0"),
        ("Quiz local scoring", "0/partial/full; answer thiếu/dư; duplicate question; unknown ID; rounding.", "Unit", "P0"),
        ("Progress merge", "Timestamp client/server; tie BestScore; pass state; timezone/UTC boundary.", "Unit", "P0"),
        ("Coins/Badges/Profile", "Parse model; loading/error/empty/success; pagination/refresh; update validation; change-password error.", "Widget", "P1"),
        ("Leaderboard", "Giữ 7 ca hiện có; thêm tie rank, >100, avatar lỗi và refresh race.", "Widget", "P1"),
        ("Lesson/Quiz screens", "Cache fallback; loading/error; select answer; submit double-tap; result; mounted guard.", "Widget", "P0/P1"),
        ("MathText/Login", "Ổn định font/render test; tách golden phụ thuộc môi trường; validation và responsive.", "Widget", "P1"),
        ("Admin screens", "CRUD form validation; optimistic state; publish validation error; delete confirmation; pagination.", "Widget", "P1/P2"),
    ]
    add_table(doc, ["Đối tượng", "Ca cần có", "Loại", "Ưu tiên"], flutter_plan, [1.45, 3.75, 0.75, 0.55], font_size=8.4)

    add_heading(doc, "6.1 Quy tắc fake/mock Flutter", 2)
    for item in [
        "Inject API/storage/clock/connectivity qua constructor hoặc interface; không gọi Dio/SQLite thật trong widget test.",
        "Fake API phải điều khiển được trạng thái pending, success, 400, 401, 500 và timeout để kiểm tra UI giữa các await.",
        "Dùng pump có giới hạn; tránh pumpAndSettle khi animation/network timer không kết thúc.",
        "Sau mỗi await dẫn đến setState, test cả trường hợp widget đã dispose để bảo vệ mounted guard.",
        "Key cho nút/input quan trọng để selector ổn định; không phụ thuộc nguyên văn toàn bộ chuỗi hiển thị.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "7. Kế hoạch API/Integration Test", 1)
    api_cases = [
        ("API-AUTH-01", "POST /api/auth/login", "Đúng/sai/inactive; schema response; không trả password hash.", "200/400/401 đúng contract", "P0"),
        ("API-AUTH-02", "JWT authorization", "Thiếu/hết hạn/sai role trên Student và Admin endpoint.", "401 hoặc 403; không lộ dữ liệu", "P0"),
        ("API-AUTH-03", "Refresh rotation", "Dùng refresh token hai lần; revoked/hết hạn.", "Chỉ lần đầu thành công", "P0"),
        ("API-QUIZ-01", "Submit lesson quiz", "Happy path và boundary pass score.", "Attempt/progress/reward đồng nhất", "P0"),
        ("API-QUIZ-02", "Submit duplicate", "Gửi cùng ClientAttemptId song song/tuần tự.", "Một attempt và một reward", "P0"),
        ("API-SYNC-01", "POST /api/sync", "Thiếu ClientAttemptId, item invalid, success, retry.", "400 ProblemDetails hoặc result đúng", "P0"),
        ("API-SYNC-02", "Partial failure", "Item 1 ghi thành công, item 2 lỗi.", "Hành vi được khóa theo contract; retry không duplicate", "P0"),
        ("API-CONTENT-01", "Student content", "Unpublished/deleted chapter, lesson, quiz.", "Không xuất hiện/404", "P0"),
        ("API-ADMIN-01", "Publish validation", "Lesson/quiz chưa đủ điều kiện.", "400/ProblemDetails; không publish", "P1"),
        ("API-LIST-01", "Pagination", "page 0/1, pageSize 0/max, ngoài range, sort/tie.", "Metadata và items ổn định", "P1"),
        ("API-PROFILE-01", "Avatar upload", "Loại file, size, empty, path traversal, URL trả về.", "Reject an toàn hoặc lưu đúng", "P1"),
        ("API-ERR-01", "Exception middleware", "Business/validation/infrastructure exception.", "Status và ProblemDetails nhất quán", "P1"),
    ]
    add_table(doc, ["ID", "Endpoint/nhóm", "Kịch bản", "Kết quả", "Mức"], api_cases, [1.0, 1.45, 2.45, 1.1, 0.5], font_size=8.1)

    add_heading(doc, "7.1 Thiết lập integration test", 2)
    for item in [
        "Dùng WebApplicationFactory<Program> và override DI cho email/generator ngoài hệ thống.",
        "Dùng PostgreSQL container/schema riêng để kiểm tra Npgsql, index unique và transaction gần production hơn InMemory provider.",
        "Tạo token Student/Admin bằng helper test với clock và secret test; tuyệt đối không dùng secret production.",
        "Seed tối thiểu bằng builder: User, Chapter, Lesson, Question, Quiz, RewardPolicy; mỗi test tự sở hữu dữ liệu.",
        "Kiểm tra cả response và side effect DB trong cùng ca, đặc biệt QuizAttempt, Progress, CoinTransaction, UserBadge, Notification.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "8. Các luồng nghiệp vụ phải kiểm thử", 1)
    add_heading(doc, "8.1 Luồng chính A - Student học online và hoàn thành quiz", 2)
    add_test_case(
        doc,
        "E2E-STU-ONLINE-01",
        "Login - học bài - quiz - nhận kết quả",
        "P0",
        "Xác nhận toàn bộ chuỗi giá trị chính của sản phẩm.",
        "Có Student active; chapter/lesson/quiz published; reward policy active; backend và PostgreSQL test env sẵn sàng.",
        [
            "Đăng nhập Student và xác nhận router vào StudentShell.",
            "Mở danh sách chương, chọn chương, mở bài học và đánh dấu đã xem nội dung.",
            "Mở quiz, trả lời đủ câu và submit đúng một lần.",
            "Đối chiếu điểm, pass state, đáp án/giải thích, coinsEarned và badge mới ở result.",
            "Mở lại dashboard, coins, badges và leaderboard để kiểm tra dữ liệu đã phản ánh.",
        ],
        "Chỉ một attempt được tạo; điểm/tiến độ/xu/huy hiệu/dashboard đồng nhất; không lộ đáp án trước submit.",
    )
    add_test_case(
        doc,
        "E2E-STU-ONLINE-02",
        "Làm lại với điểm thấp hơn",
        "P0",
        "Chứng minh pass/best score không thoái lui và retry reward đúng policy.",
        "Student đã Passed bài với BestScore cao.",
        [
            "Làm lại quiz và cố ý đạt điểm thấp hơn.",
            "Đối chiếu history và progress sau submit.",
            "Đối chiếu CoinTransaction với retry policy và daily limit.",
        ],
        "Passed giữ nguyên; BestScore không giảm; thưởng lặp lại đúng policy, không nhận first-pass bonus lần hai.",
    )

    add_heading(doc, "8.2 Luồng chính B - Offline quiz, restart và đồng bộ", 2)
    add_test_case(
        doc,
        "E2E-STU-OFFLINE-01",
        "Mất mạng - làm quiz - restart - online sync",
        "P0",
        "Chứng minh queue bền vững và đồng bộ đúng một lần.",
        "Đã cache chương/bài/câu hỏi cho đúng user; có thể bật/tắt network.",
        [
            "Tắt mạng, mở nội dung đã cache và hoàn thành quiz.",
            "Ghi lại ClientAttemptId; xác nhận attempt ở local_quiz_attempts trạng thái pending.",
            "Đóng hoàn toàn ứng dụng rồi mở lại; xác nhận queue vẫn tồn tại và thuộc đúng user.",
            "Bật mạng và kích hoạt sync; đối chiếu result, coins, pass và badges.",
            "Gửi lại cùng queue/payload hoặc mô phỏng retry sau timeout response.",
        ],
        "Server chỉ có một attempt/reward; local chỉ mark synced sau thành công; UI summary khớp server.",
    )
    add_test_case(
        doc,
        "E2E-STU-OFFLINE-02",
        "Sync lỗi và phục hồi",
        "P0",
        "Bảo vệ dữ liệu khi 400/500/timeout.",
        "Queue có nhiều attempt/progress, trong đó có một item invalid.",
        [
            "Mô phỏng 500 hoặc timeout; kiểm tra record không bị xóa, retry_count tăng, last_sync_error có giá trị an toàn.",
            "Mô phỏng 400 do payload invalid; UI hiển thị lỗi và giữ queue để sửa/retry.",
            "Sửa dữ liệu/endpoint rồi sync lại với nguyên ClientAttemptId.",
        ],
        "Không mất record; không duplicate attempt/xu/huy hiệu; trạng thái lỗi được xóa sau thành công.",
    )

    add_heading(doc, "8.3 Luồng chính C - Admin tạo và publish học liệu", 2)
    add_test_case(
        doc,
        "E2E-ADM-CONTENT-01",
        "Tạo chapter - lesson - question - quiz - publish",
        "P1",
        "Xác nhận nội dung hợp lệ đi từ Admin tới Student.",
        "Admin active; dữ liệu tên/mã không trùng.",
        [
            "Tạo chapter và lesson ở trạng thái draft.",
            "Tạo câu hỏi đủ option, đúng một correct answer và đúng lesson/chapter target.",
            "Tạo hoặc generate quiz, thêm/reorder câu hỏi.",
            "Chạy validation; thử publish trước khi đủ điều kiện và xác nhận bị chặn.",
            "Hoàn thiện nội dung, publish lesson/quiz/chapter, đăng nhập Student để xác nhận hiển thị.",
        ],
        "Draft không lộ cho Student; validation trả lỗi rõ; nội dung hợp lệ xuất hiện đúng thứ tự sau publish.",
    )

    add_heading(doc, "8.4 Luồng phụ quan trọng", 2)
    side_flows = [
        ("AUTH-REGISTER", "Gửi OTP - đăng ký - email trùng - OTP sai/hết hạn - đăng nhập lần đầu.", "P1"),
        ("AUTH-RESET", "Quên mật khẩu - OTP - reset - token/session cũ - đăng nhập bằng mật khẩu mới.", "P1"),
        ("AUTH-REFRESH", "Access token hết hạn - refresh - request tự chạy lại - refresh revoked/inactive.", "P0"),
        ("PROFILE", "Đọc/sửa hồ sơ, avatar file invalid, đổi mật khẩu, dark mode/font scale, logout.", "P1"),
        ("COINS", "Tổng xu, pagination, lịch sử quiz/badge, refresh và balanceAfter.", "P1"),
        ("BADGES", "Earned/InProgress/Locked, reconcile, legacy rule, threshold, reward/notification một lần.", "P0/P1"),
        ("LEADERBOARD", "Top 3, top 100, current user ngoài top, tie, empty/error/refresh.", "P1"),
        ("NOTIFICATIONS", "Unread count, mark one/all, admin create/delete, user isolation.", "P2"),
        ("ADMIN USERS", "Search/page/detail/history, khóa/mở tài khoản, token sau khi khóa.", "P1"),
        ("REWARD POLICY", "CRUD/active policy, giá trị biên, daily limit và hiệu lực lên attempt tiếp theo.", "P0/P1"),
        ("REPORTS", "Khoảng ngày, timezone, empty data, tổng hợp nhất quán với dữ liệu gốc.", "P2"),
        ("RESPONSIVE", "Compact/medium/wide, xoay màn hình, bàn phím che form, Android/Windows/Web.", "P2"),
    ]
    add_table(doc, ["Nhóm", "Kịch bản tối thiểu", "Mức"], side_flows, [1.35, 4.55, 0.6], font_size=8.5)

    add_heading(doc, "9. Dữ liệu test, môi trường và cô lập", 1)
    add_heading(doc, "9.1 Bộ dữ liệu chuẩn", 2)
    data_rows = [
        ("USR-STU-A", "Student active, 0 xu, chưa có progress.", "Happy path lần đầu"),
        ("USR-STU-B", "Student đã Passed, BestScore 9, có lịch sử xu/huy hiệu.", "Regression/retry"),
        ("USR-STU-C", "Student inactive, có refresh token cũ.", "Auth/active middleware"),
        ("USR-ADM-A", "Admin active.", "CRUD/publish/report"),
        ("CNT-DRAFT", "Chapter/lesson/quiz draft và thiếu điều kiện.", "Validation/visibility"),
        ("CNT-PUB", "Published lesson quiz và chapter quiz, đủ câu hỏi.", "Student online/offline"),
        ("RWD-STD", "Policy chuẩn có coins/correct, perfect, retry, first pass, daily cap.", "Reward matrix"),
        ("OFF-Q1", "Hai attempt với ClientAttemptId cố định, một progress pending.", "Sync/retry/partial failure"),
    ]
    add_table(doc, ["Mã", "Thiết lập", "Dùng cho"], data_rows, [1.05, 3.5, 1.95])

    add_heading(doc, "9.2 Quy tắc dữ liệu", 2)
    for item in [
        "UUID, email và timestamp được tạo qua TestDataBuilder; dùng UTC và clock cố định.",
        "Không phụ thuộc seed development hoặc tài khoản/mật khẩu ghi trong README khi chạy CI.",
        "Mỗi integration test dùng transaction/schema riêng; cleanup không được trỏ tới database development/production.",
        "Fixture quiz phải khai báo rõ target LessonId xor ChapterId và IsPublished/IsDeleted.",
        "Dữ liệu thưởng phải kiểm tra cả tổng User.Coins và ledger CoinTransaction/BalanceAfter.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "10. Roadmap triển khai", 1)
    roadmap = [
        ("Phase 0 - 1 ngày", "Ổn định baseline", "Tách Flutter suite, xử lý timeout/flaky MathText/Login; ghi test command và report chuẩn.", "Suite có kết quả lặp lại 3 lần"),
        ("Phase 1 - 2 đến 3 ngày", "Khóa P0 backend", "Mở rộng Auth, QuizScoring, Reward, Progress, Badge; thêm invariant/idempotency/concurrency.", "P0 unit đạt 100%"),
        ("Phase 2 - 2 đến 3 ngày", "Khóa Flutter offline/UI", "Test OfflineSyncService/LocalDbService, Coins/Badges/Profile và quiz double-submit.", "P0/P1 Flutter đạt"),
        ("Phase 3 - 3 đến 4 ngày", "API integration", "WebApplicationFactory + PostgreSQL; JWT/role, submit, sync, publish, pagination, ProblemDetails.", "Integration P0 đạt 100%"),
        ("Phase 4 - 2 ngày", "E2E và release gate", "Android offline/restart/sync; Admin publish; smoke Windows/Web; report coverage.", "Không blocker/critical"),
    ]
    add_table(doc, ["Giai đoạn", "Trọng tâm", "Công việc", "Điều kiện hoàn tất"], roadmap, [1.2, 1.25, 2.9, 1.15], font_size=8.5)

    add_heading(doc, "10.1 Ước lượng danh mục test mục tiêu", 2)
    add_table(
        doc,
        ["Suite", "Hiện có quan sát từ source", "Mục tiêu gần", "Ghi chú"],
        [
            ("Backend unit", "27 test", "70-90 test", "Tăng branch P0/P1, không chạy theo số lượng đơn thuần."),
            ("Flutter unit", "Router, scoring, merge, offline summary, responsive.", "45-60 test", "Bổ sung ApiClient/DB/sync/model."),
            ("Flutter widget", "Login, MathText, responsive, leaderboard.", "45-65 test", "Bổ sung màn hình Student/Admin quan trọng."),
            ("API integration", "Chưa thấy suite chuyên biệt.", "25-35 test", "Ưu tiên JWT, quiz, sync, publish, pagination."),
            ("E2E", "Login Student/Admin và responsive scaffold.", "8-12 hành trình", "Ít nhưng tập trung P0/P1."),
        ],
        [1.25, 2.4, 1.25, 1.6],
    )

    add_heading(doc, "11. Quality gate và tiêu chí hoàn tất", 1)
    gates = [
        ("Pull request", "Build backend/Flutter phần liên quan; unit test thay đổi đạt; không test flaky mới; không giảm coverage file bị sửa > 2 điểm %.", "Bắt buộc"),
        ("Merge nhánh chính", "Tất cả P0/P1 unit + API integration đạt; flutter analyze không có issue mới; secret/debug body log không xuất hiện.", "Bắt buộc"),
        ("Release candidate", "E2E online + offline/restart/sync đạt; Android smoke; migration apply trên DB disposable; 0 blocker/critical.", "Bắt buộc"),
        ("Release", "PO/QA chấp nhận known issues P2/P3; report test/coverage lưu cùng build; rollback/config đã xác nhận.", "Bắt buộc"),
    ]
    add_table(doc, ["Cổng", "Điều kiện", "Trạng thái"], gates, [1.4, 4.35, 0.75])

    add_heading(doc, "11.1 Definition of Done cho một test task", 2)
    for item in [
        "Test name thể hiện Given/When/Then và liên kết ID invariant/requirement nếu là P0/P1.",
        "Có success, validation, authorization/not-found và side-effect assertion phù hợp.",
        "Test chạy độc lập, không phụ thuộc thứ tự, timezone, network thật hoặc dữ liệu từ test khác.",
        "Không để token/password/đáp án trong log; fixture không chứa secret production.",
        "Đã chạy suite liên quan ít nhất 3 lần nếu vừa sửa flaky/async/widget test.",
        "TASKS.md và tài liệu test được cập nhật khi contract hoặc workflow thay đổi.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "12. Lệnh chạy và tổ chức CI", 1)
    add_table(
        doc,
        ["Mục đích", "Lệnh", "Kỳ vọng"],
        [
            ("Backend build", "cd backend; dotnet build MathIBook.sln", "0 error; warning mới phải triage"),
            ("Backend test", "cd backend; dotnet test MathIBook.sln", "P0/P1 đạt; xuất TRX/coverage trong CI"),
            ("Flutter analyze", "cd frontend; flutter analyze", "Không tăng issue; tiến tới 0 issue"),
            ("Flutter unit/widget", "cd frontend; flutter test", "Tách shard/tag nếu vượt timeout"),
            ("Flutter integration", "cd frontend; flutter test integration_test", "Chạy với backend test env"),
            ("Format", "dart format lib test integration_test", "Không có diff format ngoài ý muốn"),
        ],
        [1.45, 3.0, 2.05],
        font_size=8.8,
    )
    add_body(doc, "Đề xuất pipeline: restore/cache -> backend build + unit -> Flutter analyze + unit/widget theo shard -> API integration với PostgreSQL -> build artifact -> E2E theo lịch hoặc release candidate.")

    add_heading(doc, "13. Theo dõi defect và báo cáo", 1)
    add_table(
        doc,
        ["Severity", "Định nghĩa", "Ví dụ", "SLA xử lý đề xuất"],
        [
            ("Blocker", "Không thể test/release hoặc mất dữ liệu diện rộng.", "Không login/submit/sync được.", "Xử lý ngay; dừng release"),
            ("Critical", "Sai quyền, duplicate reward, mất queue, lộ dữ liệu nhạy cảm.", "Cộng xu hai lần khi retry.", "Trong ngày; thêm regression test"),
            ("Major", "Sai chức năng P1, có workaround.", "Badge không reconcile, publish sai.", "Trước release"),
            ("Minor", "UX/cosmetic hoặc edge P2/P3.", "Empty state/căn chỉnh.", "Theo backlog"),
        ],
        [1.0, 2.35, 2.05, 1.1],
    )
    add_body(doc, "Mỗi defect cần: môi trường/build, user/role, dữ liệu seed, bước tái hiện, actual/expected, response/status, ClientAttemptId hoặc correlation ID đã che thông tin nhạy cảm, ảnh/video khi là UI, và test regression được thêm sau khi sửa.")

    add_heading(doc, "14. Ma trận truy vết rút gọn", 1)
    trace = [
        ("REQ-01 Học và quiz", "INV-02,03,06", "QuizScoring/Domain", "API-QUIZ-01", "E2E-STU-ONLINE-01/02"),
        ("REQ-02 Reward", "INV-01,04,05", "Reward/Badge", "API-QUIZ-02", "Coins/Badges/Leaderboard"),
        ("REQ-03 Offline", "INV-01,08,09", "OfflineSync/LocalDb", "API-SYNC-01/02", "E2E-STU-OFFLINE-01/02"),
        ("REQ-04 Auth/role", "INV-07,10", "Auth/Router/ApiClient", "API-AUTH-01/02/03", "Student/Admin login"),
        ("REQ-05 Content admin", "INV-06", "ContentValidation", "API-CONTENT/ADMIN", "E2E-ADM-CONTENT-01"),
        ("REQ-06 Profile/notify", "User isolation", "Profile/model/widget", "API-PROFILE", "Side-flow smoke"),
    ]
    add_table(doc, ["Yêu cầu", "Invariant", "Unit", "Integration", "E2E/Widget"], trace, [1.35, 1.0, 1.4, 1.45, 1.3], font_size=8.3)

    add_heading(doc, "15. Checklist trước khi release", 1)
    release_checks = [
        "Backend P0/P1 unit test và API integration đạt 100%.",
        "Flutter suite có kết quả cuối, không timeout; test flaky đã được cách ly có ticket và owner.",
        "Online quiz và retry điểm thấp hơn không làm mất pass/best score.",
        "Offline quiz qua restart rồi sync không duplicate attempt/xu/huy hiệu.",
        "Student/Admin JWT và tài khoản inactive trả đúng 401/403.",
        "Draft/unpublished/deleted content không xuất hiện cho Student.",
        "Coins ledger, User.Coins, badges, notification và leaderboard nhất quán.",
        "Không log password/token/body đáp án; cấu hình test không chứa secret production.",
        "Migration chạy trên database disposable và rollback/recovery đã được xác nhận.",
        "Test report, coverage, known issues và build artifact được lưu cùng release candidate.",
    ]
    for item in release_checks:
        add_list_item(doc, item)

    add_callout(
        doc,
        "Điểm chốt:",
        "Không coi luồng offline là đạt chỉ vì request sync trả 200. Phải đối chiếu cả queue local, attempt server, reward ledger, pass/best score, badge và hành vi retry cùng ClientAttemptId.",
        RED,
    )

    add_heading(doc, "Phụ lục A - Quy ước đặt tên test", 1)
    add_table(
        doc,
        ["Loại", "Mẫu", "Ví dụ"],
        [
            ("C# unit", "Method_State_Expected", "ScoreQuizAsync_DuplicateClientAttempt_ReturnsExistingWithoutReward"),
            ("Dart unit", "given ... when ... then ...", "given failed queue when retry succeeds then marks attempts synced"),
            ("Widget", "screen + state + behavior", "CoinsScreen error state shows retry and reloads"),
            ("Integration", "endpoint + condition + contract", "POST sync missing clientAttemptId returns 400 ProblemDetails"),
            ("E2E", "role + journey + outcome", "Student offline quiz restart sync awards once"),
        ],
        [1.0, 2.4, 3.1],
    )

    add_heading(doc, "Phụ lục B - Nguồn đối chiếu trong repository", 1)
    for item in [
        "Backend: backend/src/MathIBook.Application/Services và backend/src/MathIBook.Api/Controllers.",
        "Backend tests: backend/tests/MathIBook.UnitTests.",
        "Flutter core: frontend/lib/core/network, core/storage, core/sync, core/router.",
        "Flutter features: frontend/lib/features/student và frontend/lib/features/admin.",
        "Flutter tests: frontend/test và frontend/integration_test.",
        "Quy ước dự án: README.md, AGENTS.md và TASKS.md.",
    ]:
        add_list_item(doc, item)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Kế hoạch Test và Unit Test - Math IBook"
    doc.core_properties.subject = "Luồng chính và các luồng phụ quan trọng"
    doc.core_properties.author = "Math IBook Project Team"
    doc.core_properties.keywords = "Test Plan, Unit Test, Integration Test, Flutter, ASP.NET Core, Offline Sync"
    doc.core_properties.comments = "Generated from the repository baseline on 20/07/2026."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
